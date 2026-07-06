"""Gerador de documentos .docx institucionais para RGF Despesa com Pessoal — TCE-AL."""

import io
import pathlib
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from .rgf_limites import Situacao, get_limite_maximo, label_periodo_completo
from .siconfi_rgf_client import ResultadoConsultaRGF

# Caminho base dos modelos Word — templates/alertas/ dentro do repositório
_DIR_MODELOS = pathlib.Path(__file__).parent.parent / "templates" / "alertas"

_AZUL_TCE = RGBColor(0x1F, 0x4E, 0x79)

_LABEL_SITUACAO: dict[Situacao, str] = {
    Situacao.NORMAL:     "Dentro do Limite",
    Situacao.ALERTA:     "Alerta (Art. 59, §1º, II)",
    Situacao.PRUDENCIAL: "Prudencial (Art. 22)",
    Situacao.MAXIMO:     "Excede Limite Máximo (Arts. 19/20)",
    Situacao.SEM_DADOS:  "Sem dados",
}

_ESFERA_LABEL = {"M": "Municipal", "E": "Estadual"}
_PODER_LABEL  = {"E": "Executivo", "L": "Legislativo"}


def _fmt_pct(valor: float) -> str:
    return f"{valor:.2f}%".replace(".", ",")


def _cabecalho(doc: Document, titulo: str, subtitulo: str) -> None:
    """Adiciona cabeçalho institucional TCE-AL ao documento."""
    h1 = doc.add_heading(level=1)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h1.add_run("TRIBUNAL DE CONTAS DO ESTADO DE ALAGOAS — TCE-AL")
    run.font.color.rgb = _AZUL_TCE
    run.font.size = Pt(13)

    h2 = doc.add_heading(level=2)
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = h2.add_run("Diretoria de Coordenação de Técnicos — DCT")
    run2.font.color.rgb = _AZUL_TCE
    run2.font.size = Pt(11)

    doc.add_paragraph()

    ht = doc.add_heading(level=2)
    ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ht.add_run(titulo).font.size = Pt(12)

    ps = doc.add_paragraph()
    ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ps.add_run(subtitulo).font.size = Pt(10)

    doc.add_paragraph()


def _rodape_legal(doc: Document) -> None:
    """Adiciona nota de rodapé com base legal e data de geração."""
    doc.add_paragraph()
    p1 = doc.add_paragraph()
    r1 = p1.add_run(
        "Base legal: Lei Complementar nº 101/2000 (LRF) — Arts. 19, 20, 22 e 59, §1º, II. "
        "Resolução Normativa TCE-AL nº 5/2024."
    )
    r1.font.size = Pt(8)
    r1.italic = True

    p2 = doc.add_paragraph()
    r2 = p2.add_run(
        f"Documento gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M:%S')} "
        "pela ferramenta SICONFI-TCE-AL (DCT)."
    )
    r2.font.size = Pt(8)
    r2.italic = True


def _subtitulo_padrao(resultado: ResultadoConsultaRGF, nr_quadrimestre: int, nr_semestre: int) -> str:
    label_periodo = label_periodo_completo(nr_quadrimestre, nr_semestre)
    esfera = _ESFERA_LABEL.get(resultado.esfera, resultado.esfera)
    poder = _PODER_LABEL.get(resultado.poder, resultado.poder)
    return (
        f"Despesa Total com Pessoal — {esfera} / {poder} · "
        f"{label_periodo} / {resultado.exercicio}"
    )


# ---------------------------------------------------------------------------
# Labels de situação conforme o modelo institucional
# ---------------------------------------------------------------------------
_LABEL_SITUACAO_MODELO: dict[Situacao, str] = {
    Situacao.ALERTA:     "Acima do Limite de Alerta",
    Situacao.PRUDENCIAL: "Acima do Limite Prudencial",
    Situacao.MAXIMO:     "Acima do Limite Máximo",
}

_ORDINAL_BR = {1: "1º", 2: "2º", 3: "3º"}


def _label_unidade_gestora(poder: str) -> str:
    """Retorna o prefixo da Unidade Gestora conforme o poder."""
    return "Prefeitura Municipal" if poder == "E" else "Câmara Municipal"


def _remover_linhas_dados_tabela(table) -> None:
    """Remove todas as linhas de dados da tabela, mantendo apenas o cabeçalho (linha 0)."""
    tbl = table._tbl
    for tr in list(tbl.findall(qn("w:tr")))[1:]:
        tbl.remove(tr)


def _adicionar_linha_tabela_modelo(
    table,
    unidade_gestora: str,
    municipio: str,
    percentual: str,
    situacao: str,
    font_size_twips: int = 139700,
) -> None:
    """Adiciona uma linha de dados à tabela com formatação idêntica ao modelo."""
    row = table.add_row()
    valores = [unidade_gestora, municipio, percentual, situacao]
    for cell, texto in zip(row.cells, valores):
        # Limpa runs existentes e define o texto preservando tamanho de fonte
        para = cell.paragraphs[0]
        for run in para.runs:
            run.text = ""
        if para.runs:
            run = para.runs[0]
        else:
            run = para.add_run()
        run.text = texto
        run.font.size = font_size_twips
        run.bold = False


def gerar_alerta_docx(
    resultado: ResultadoConsultaRGF,
    nr_quadrimestre: int,
    nr_semestre: int,
    data_extracao: str | None = None,
) -> bytes:
    """Gera Termo de Alerta .docx usando o modelo institucional como base.

    Abre o documento template correspondente ao poder (Executivo/Legislativo),
    substitui o período e a data de extração nos parágrafos corretos e
    popula a tabela com os entes críticos do resultado.

    Args:
        resultado: Resultado da consulta RGF.
        nr_quadrimestre: Número do quadrimestre consultado.
        nr_semestre: Número do semestre consultado.
        data_extracao: Data de extração dos dados no formato DD/MM/AAAA.
                       Se None, usa a data atual.

    Returns:
        Bytes do documento .docx gerado.
    """
    # Seleciona o template conforme o poder
    nome_arquivo = (
        "Alerta de Despesa de Pessoal - Poder Executivo.docx"
        if resultado.poder == "E"
        else "Alerta de Despesa de Pessoal - Poder Legislativo.docx"
    )
    template_path = _DIR_MODELOS / nome_arquivo
    doc = Document(str(template_path))

    # ── Período a substituir no parágrafo 4 ──────────────────────────────────
    # O parágrafo contém runs separados para cada parte do período:
    #   run[3] = número do quadrimestre (ex: "3")
    #   run[4] = "º Quadrimestre"
    #   run[5] = " "
    #   run[6] = "e Xº Semestre " (contém o ordinal do semestre)
    #   run[7] = "de AAAA" (contém o ano)
    para_intro = doc.paragraphs[4]
    runs = para_intro.runs

    nr_quad_ord = _ORDINAL_BR.get(nr_quadrimestre, f"{nr_quadrimestre}º")
    nr_sem_ord  = _ORDINAL_BR.get(nr_semestre,    f"{nr_semestre}º")

    # Substitui número do quadrimestre (run 3 é apenas o dígito)
    if len(runs) > 3:
        runs[3].text = nr_quad_ord[:-1]  # só o número, sem o "º" (run 4 tem "º Quadrimestre")
    # Substitui ordinal do semestre no run 6
    if len(runs) > 6:
        runs[6].text = f"e {nr_sem_ord} Semestre "
    # Substitui ano no run 7
    if len(runs) > 7:
        runs[7].text = f"de {resultado.exercicio}"

    # ── Data de extração no parágrafo 10 ─────────────────────────────────────
    # Executivo: run[0] = texto completo até data, run[1] = data
    # Legislativo: run[0..2] = texto, run[3] = data
    data_str = data_extracao or datetime.now().strftime("%d/%m/%Y")
    para_fonte = doc.paragraphs[10]
    if resultado.poder == "E":
        # Executivo: run[1] contém apenas a data
        if len(para_fonte.runs) > 1:
            para_fonte.runs[1].text = data_str
    else:
        # Legislativo: run[3] contém apenas a data
        if len(para_fonte.runs) > 3:
            para_fonte.runs[3].text = data_str
        elif len(para_fonte.runs) > 1:
            para_fonte.runs[-1].text = data_str

    # ── Título da tabela (parágrafo 9) ───────────────────────────────────────
    # Executivo: run[3] = ' — Poder Executivo'  (4 runs total)
    # Legislativo: run[3] = ' — Poder ', run[4] = 'Legislativo'  (5 runs total)
    # Como cada template já tem o nome correto do poder, não é necessário alterar
    # se o poder selecionado corresponde ao template carregado. Mas por segurança,
    # mantemos a lógica para garantir consistência.
    para_titulo_tabela = doc.paragraphs[9]
    poder_label = "Executivo" if resultado.poder == "E" else "Legislativo"
    if resultado.poder == "E":
        # run[3] = ' — Poder Executivo'
        if len(para_titulo_tabela.runs) > 3:
            para_titulo_tabela.runs[3].text = f" — Poder {poder_label}"
    else:
        # run[3] = ' — Poder ', run[4] = 'Legislativo'
        if len(para_titulo_tabela.runs) > 4:
            para_titulo_tabela.runs[3].text = " — Poder "
            para_titulo_tabela.runs[4].text = poder_label
        elif len(para_titulo_tabela.runs) > 3:
            para_titulo_tabela.runs[3].text = f" — Poder {poder_label}"

    # ── Popula a tabela com os entes críticos ─────────────────────────────────
    criticos = [
        e for e in resultado.entes_com_dados
        if e.situacao in (Situacao.ALERTA, Situacao.PRUDENCIAL, Situacao.MAXIMO)
    ]

    tabela = doc.tables[0]
    _remover_linhas_dados_tabela(tabela)

    unidade_gestora_label = _label_unidade_gestora(resultado.poder)

    for ente in sorted(criticos, key=lambda e: (-e.percentual_dtp, e.nome)):
        situacao_label = _LABEL_SITUACAO_MODELO.get(ente.situacao, ente.situacao.value)
        percentual_fmt = f"     {_fmt_pct(ente.percentual_dtp).replace('%', '')} "
        _adicionar_linha_tabela_modelo(
            tabela,
            unidade_gestora=f"{unidade_gestora_label} ",
            municipio=ente.nome,
            percentual=percentual_fmt,
            situacao=situacao_label,
        )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def gerar_relatorio_docx(
    resultado: ResultadoConsultaRGF,
    nr_quadrimestre: int,
    nr_semestre: int,
) -> bytes:
    """Gera Relatório de Gestão Fiscal .docx com todos os entes e análise LC 178/2021.

    Args:
        resultado: Resultado da consulta RGF.
        nr_quadrimestre: Número do quadrimestre consultado.
        nr_semestre: Número do semestre consultado.

    Returns:
        Bytes do documento .docx.
    """
    doc = Document()
    _cabecalho(doc, "RELATÓRIO DE GESTÃO FISCAL — DESPESA COM PESSOAL", _subtitulo_padrao(resultado, nr_quadrimestre, nr_semestre))

    # Seção 1 — Resumo por situação
    doc.add_heading("1. Resumo por Situação", level=2)

    contagem: dict[Situacao, int] = {s: 0 for s in Situacao}
    for ente in resultado.entes_com_dados:
        contagem[ente.situacao] += 1

    t_resumo = doc.add_table(rows=1, cols=2)
    t_resumo.style = "Table Grid"
    hdr = t_resumo.rows[0].cells
    hdr[0].text, hdr[1].text = "Situação", "Quantidade de Entes"
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True

    for sit in (Situacao.MAXIMO, Situacao.PRUDENCIAL, Situacao.ALERTA, Situacao.NORMAL):
        row = t_resumo.add_row().cells
        row[0].text = _LABEL_SITUACAO.get(sit, sit.value)
        row[1].text = str(contagem[sit])

    row_sd = t_resumo.add_row().cells
    row_sd[0].text = "Sem dados (inadimplentes)"
    row_sd[1].text = str(len(resultado.entes_sem_dados))

    doc.add_paragraph()

    # Seção 2 — Tabela completa
    doc.add_heading("2. Situação de Todos os Entes", level=2)

    if resultado.entes_com_dados:
        t_entes = doc.add_table(rows=1, cols=4)
        t_entes.style = "Table Grid"
        hdr2 = t_entes.rows[0].cells
        for i, texto in enumerate(["Ente Fiscal", "% DTP/RCL", "Período", "Situação"]):
            hdr2[i].text = texto
            hdr2[i].paragraphs[0].runs[0].bold = True

        for ente in resultado.entes_com_dados:
            row = t_entes.add_row().cells
            row[0].text = ente.nome
            row[1].text = _fmt_pct(ente.percentual_dtp)
            row[2].text = f"{ente.nr_periodo}º {'Quad.' if ente.periodo_tipo == 'Q' else 'Sem.'}"
            row[3].text = _LABEL_SITUACAO.get(ente.situacao, ente.situacao.value)
    else:
        doc.add_paragraph("Nenhum ente com dados disponíveis no período consultado.")

    doc.add_paragraph()

    # Seção 3 — Inadimplentes
    doc.add_heading("3. Entes sem Dados no SICONFI (Inadimplentes)", level=2)

    if resultado.entes_sem_dados:
        doc.add_paragraph(
            f"Os {len(resultado.entes_sem_dados)} entes a seguir não publicaram o "
            "Demonstrativo da Despesa com Pessoal (RGF Anexo 01) no período consultado:"
        )
        for nome in resultado.entes_sem_dados:
            doc.add_paragraph(f"• {nome}")
    else:
        doc.add_paragraph("Todos os entes consultados publicaram os dados no SICONFI.")

    doc.add_paragraph()

    # Seção 4 — Análise LC 178/2021
    doc.add_heading("4. Análise LC 178/2021 — Regime de Recuperação Fiscal", level=2)

    acima_limite = [e for e in resultado.entes_com_dados if e.situacao == Situacao.MAXIMO]
    limite_max = get_limite_maximo(resultado.esfera, resultado.poder)

    if acima_limite:
        doc.add_paragraph(
            "Os entes abaixo ultrapassaram o limite máximo de despesa com pessoal. "
            "Nos termos da Lei Complementar nº 178/2021 (PATF), esses entes podem estar "
            "sujeitos a restrições na obtenção de crédito e transferências voluntárias:"
        )
        for ente in sorted(acima_limite, key=lambda e: -e.percentual_dtp):
            excesso = ente.percentual_dtp - limite_max
            doc.add_paragraph(
                f"• {ente.nome}: {_fmt_pct(ente.percentual_dtp)} "
                f"(excede o limite máximo em {_fmt_pct(excesso)})"
            )
    else:
        doc.add_paragraph(
            "Nenhum ente ultrapassou o limite máximo de despesa com pessoal no período "
            "consultado. Não se aplica o Regime de Recuperação Fiscal previsto na LC 178/2021."
        )

    _rodape_legal(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

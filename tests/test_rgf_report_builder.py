"""Testes para src/rgf_report_builder.py — geração do Termo de Alerta."""

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.rgf_report_builder import gerar_alerta_docx
from src.rgf_limites import Situacao
from src.siconfi_rgf_client import ResultadoConsultaRGF, ResultadoEnteRGF


def _resultado(entes_sem_dados: list[str]) -> ResultadoConsultaRGF:
    ente_critico = ResultadoEnteRGF(
        id_ente="2700102",
        nome="Arapiraca",
        percentual_dtp=55.0,
        situacao=Situacao.ALERTA,
        periodo_tipo="Q",
        nr_periodo=1,
    )
    return ResultadoConsultaRGF(
        entes_com_dados=[ente_critico],
        entes_sem_dados=entes_sem_dados,
        exercicio=2025,
        esfera="M",
        poder="E",
    )


def _paragraphs_texto(doc_bytes: bytes) -> list[str]:
    doc = Document(io.BytesIO(doc_bytes))
    return [p.text for p in doc.paragraphs]


def test_alerta_com_entes_sem_dados_inclui_texto_e_lista_numerada():
    resultado = _resultado(["Maceió", "Palmeira dos Índios"])

    doc_bytes = gerar_alerta_docx(
        resultado,
        nr_quadrimestre=1,
        nr_semestre=1,
        data_extracao="05/07/2026",
        tipo_periodo="A",
    )

    textos = _paragraphs_texto(doc_bytes)

    assert any(
        "não haviam enviado ao SICONFI" in t and "05/07/2026" in t
        for t in textos
    )
    assert "1. Maceió" in textos
    assert "2. Palmeira dos Índios" in textos

    # A seção deve ficar após a tabela ("Fonte: ...") e antes das assinaturas
    # ("Tribunal de Contas do Estado de Alagoas").
    idx_fonte = next(i for i, t in enumerate(textos) if t.startswith("Fonte:"))
    idx_assinatura = next(
        i for i, t in enumerate(textos)
        if t == "Tribunal de Contas do Estado de Alagoas"
    )
    idx_intro = next(
        i for i, t in enumerate(textos) if "não haviam enviado ao SICONFI" in t
    )
    idx_item1 = textos.index("1. Maceió")
    idx_item2 = textos.index("2. Palmeira dos Índios")
    assert idx_fonte < idx_intro < idx_item1 < idx_item2 < idx_assinatura

    # Formatação deve seguir o padrão do corpo do documento: Times New Roman,
    # sem negrito/itálico, justificado.
    doc = Document(io.BytesIO(doc_bytes))
    for texto_alvo in (
        "1. Maceió",
        "2. Palmeira dos Índios",
    ):
        paragrafo = next(p for p in doc.paragraphs if p.text == texto_alvo)
        assert paragrafo.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in paragrafo.runs:
            assert run.font.name == "Times New Roman"
            assert not run.bold
            assert not run.italic


def test_alerta_sem_entes_sem_dados_nao_inclui_secao():
    resultado = _resultado([])

    doc_bytes = gerar_alerta_docx(
        resultado,
        nr_quadrimestre=1,
        nr_semestre=1,
        data_extracao="05/07/2026",
        tipo_periodo="A",
    )

    textos = _paragraphs_texto(doc_bytes)

    assert not any("não haviam enviado ao SICONFI" in t for t in textos)
